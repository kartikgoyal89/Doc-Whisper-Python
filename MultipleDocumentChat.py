import pdfplumber
import docx
import pytesseract
import pandas as pd
from PIL import Image
import requests
import io
from langchain_huggingface import HuggingFaceEmbeddings
from fastapi import FastAPI, HTTPException,Body
from pydantic import BaseModel
from langchain_text_splitters import CharacterTextSplitter,RecursiveCharacterTextSplitter
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from fastapi.middleware.cors import CORSMiddleware
import os
from dotenv import load_dotenv
import faiss
import uvicorn
import numpy as np
from langchain_groq import ChatGroq
import tempfile



load_dotenv()
groq_api_key = os.getenv("GROQ_API_KEY")

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


embedding_model = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")

def download_file(url):
    response = requests.get(url)
    return io.BytesIO(response.content)


def extract_text_and_tables_from_pdf(file_url):
    """
    Extracts text and tables from a PDF file while maintaining their order.
    Returns the extracted content as a string.
    
    :param file_bytes: Bytes of the PDF file.
    :return: Extracted text and tables as a string.
    """
    extracted_content = ""  # Store extracted content as a string

    with pdfplumber.open(file_url) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            extracted_content += f"\n--- Page {page_num} ---\n"

            elements = []  # Store text and table elements

            # Extract text as individual words with their positions
            words = page.extract_words()
            for word in words:
                elements.append((word['top'], word['text'], "text"))

            # Extract tables with positions
            tables = page.extract_tables()
            if tables:
                for table in tables:
                    if not table or not table[0]:  # Handle empty tables
                        continue

                    # Replace None values with empty strings ("")
                    table = [[cell if cell is not None else "" for cell in row] for row in table]

                    # Convert table to Markdown format
                    table_str = "| " + " | ".join(table[0]) + " |\n"  # Header row
                    table_str += "| " + " | ".join(["-" * len(col) for col in table[0]]) + " |\n"  # Separator row
                    for row in table[1:]:
                        table_str += "| " + " | ".join(row) + " |\n"

                    # Wrap the table in markdown-style code block for easy parsing
                    formatted_table = f"\n```\n{table_str}\n```\n"
                    elements.append((page.bbox[1], formatted_table, "table"))

            # Sort elements based on their position (top to bottom)
            elements.sort(key=lambda x: x[0])

            # Append content in the correct order
            for _, content, elem_type in elements:
                if elem_type == "text":
                    extracted_content += content + " "  # Append space for readability
                elif elem_type == "table":
                    extracted_content += "\n" + content + "\n"  # Properly formatted markdown table

    return extracted_content  # Return extracted content as a string



def extract_text_and_tables_from_docx(file_url):
    """
    Extracts text and tables from a DOCX file, maintaining order.
    
    :param file_bytes: Bytes of the DOCX file.
    :return: Combined text with formatted tables in Markdown format.
    """
    doc = docx.Document(file_url)
    
    elements = []  # Stores text and tables with positions

    # Extract paragraphs (text)
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            elements.append((idx, text, "text"))  # Store with index to maintain order

    # Extract tables
    for idx, table in enumerate(doc.tables, start=len(doc.paragraphs)):  # Ensure order after paragraphs
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() if cell.text else "" for cell in row.cells]  # Handle empty cells
            table_data.append(row_data)

        # Convert table to Markdown format
        if table_data:
            table_str = "| " + " | ".join(table_data[0]) + " |\n"
            table_str += "| " + " | ".join(["-" * len(col) for col in table_data[0]]) + " |\n"
            for row in table_data[1:]:
                table_str += "| " + " | ".join(row) + " |\n"

            formatted_table = f"\n```\n{table_str}\n```\n"
            elements.append((idx, formatted_table, "table"))

    # Sort elements by their order (ensuring text and tables appear as in the DOCX file)
    elements.sort(key=lambda x: x[0])

    # Combine everything into a single formatted string
    extracted_content = " ".join(content for _, content, _ in elements)

    return extracted_content


def extract_text_and_tables_from_xlsx(file_url):
    df = pd.read_excel(file_url, sheet_name=None)
    tables = {sheet: df[sheet].to_json() for sheet in df}
    text = ""  # XLSX files typically don't have freeform text
    return tables


def extract_text_from_txt(file_like):
    """
    Extracts text from a plain .txt file.

    :param file_like: BytesIO or file-like object (or string path or URL).
    :return: Extracted text as a string.
    """

    if isinstance(file_like, io.BytesIO):
        try:
            return file_like.getvalue().decode('utf-8').strip()
        except UnicodeDecodeError:
            return file_like.getvalue().decode('latin1').strip()

    elif isinstance(file_like, str):
        if os.path.exists(file_like):  # Local file
            try:
                with open(file_like, 'r', encoding='utf-8') as f:
                    return f.read().strip()
            except UnicodeDecodeError:
                with open(file_like, 'r', encoding='latin1') as f:
                    return f.read().strip()
        else:  # Assume it's a URL
            try:
                response = requests.get(file_like)
                response.raise_for_status()
                try:
                    return response.content.decode('utf-8').strip()
                except UnicodeDecodeError:
                    return response.content.decode('latin1').strip()
            except Exception as e:
                print(f"Error fetching TXT file: {e}")
                return ""

    else:
        print("Unsupported file format for TXT processing.")
        return ""

"""
def extract_from_excel(file_url):
    # Convert file bytes into a file-like object
    excel_file = io.BytesIO(file_url)

    # Read all sheets in the Excel file
    df_sheets = pd.read_excel(excel_file, sheet_name=None)

    text_output = ""

    for sheet_name, df in df_sheets.items():
        text_output += f"\n--- Sheet: {sheet_name} ---\n"

        if df.empty:
            continue

        # Replace NaN with empty strings
        df = df.fillna("")

        # Convert header row
        header = "| " + " | ".join(str(col) for col in df.columns) + " |"
        separator = "| " + " | ".join("-" * len(str(col)) for col in df.columns) + " |"

        table_rows = [header, separator]

        # Convert each row into markdown format
        for _, row in df.iterrows():
            row_data = "| " + " | ".join(str(cell) for cell in row) + " |"
            table_rows.append(row_data)

        # Join and format
        markdown_table = "\n".join(table_rows)
        text_output += f"\n```\n{markdown_table}\n```\n"

    return text_output, []  # Return text, empty tables (if needed)
"""


def extract_from_excel_or_csv(file_url):
    text_output = ""

    # Download the file from URL to a temporary file
    response = requests.get(file_url)
    response.raise_for_status()  # Raise error if download fails

    # Guess file extension from URL
    _, file_ext = os.path.splitext(file_url)

    # Create a temporary file with the correct extension
    with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as tmp_file:
        tmp_file.write(response.content)
        tmp_file_path = tmp_file.name

    # Process the file
    if file_ext.lower() == ".csv":
        df = pd.read_csv(tmp_file_path)
        df_sheets = {"Sheet1": df}
    else:
        df_sheets = pd.read_excel(tmp_file_path, sheet_name=None)

    for sheet_name, df in df_sheets.items():
        text_output += f"\n--- Sheet: {sheet_name} ---\n"

        if df.empty:
            continue

        df = df.fillna("")

        header = "| " + " | ".join(str(col) for col in df.columns) + " |"
        separator = "| " + " | ".join("-" * len(str(col)) for col in df.columns) + " |"
        table_rows = [header, separator]

        for _, row in df.iterrows():
            row_data = "| " + " | ".join(str(cell) for cell in row) + " |"
            table_rows.append(row_data)

        markdown_table = "\n".join(table_rows)
        text_output += f"\n```\n{markdown_table}\n```\n"

    # Clean up the downloaded temp file
    os.remove(tmp_file_path)

    return text_output, []



pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"


def extract_text_from_image(file_url):
    image = Image.open(file_url)
    text = pytesseract.image_to_string(image)
    return text, []


class FileRequest(BaseModel):
    fileUrl: str

def process_document(fileUrl):
    file_url = fileUrl
    file_bytes = download_file(file_url)
    text = ""

    # Try PDF
    try:
        text = extract_text_and_tables_from_pdf(file_bytes)
        if text:  # If PDF returned valid content, return early
            return text
    except Exception as e:
        pass

    # Try DOCX
    try:
        text = extract_text_and_tables_from_docx(file_bytes)
        if text:
            return text
    except Exception as e:
        pass

    # Try TXT
    try:
        text = extract_text_from_txt(file_bytes)
        if text:
            return text
    except Exception as e:
        pass

    # Try XLSX
    try:
        text, tables = extract_text_and_tables_from_xlsx(file_bytes)
        if text or tables:
            return text
    except Exception as e:
        pass

    # Try Image
    try:
        text, tables = extract_text_from_image(file_bytes)
        return text
    except Exception as e:
        pass

    # If all fail
    return ""


class EmbeddingRequest(BaseModel):
    fileUrl: str

@app.post("/generate-embeddings")
def generate_embeddings(request: EmbeddingRequest):
    file_url = request.fileUrl
    try:
        # text_splitter = RecursiveCharacterTextSplitter(separator=" ",chunk_size=1000, chunk_overlap=200)
        text_splitter = CharacterTextSplitter(
            separator=" ",  # ✅ Ensure splitting works well for PDFs
            chunk_size=800,
            chunk_overlap=50
        )
        processed_data = process_document(file_url)
        # text_chunks = text_splitter.split_text(processed_data)
        text_chunks = text_splitter.split_text(processed_data)


        embeddings = embedding_model.embed_documents(text_chunks)
        

        return { "embeddings": embeddings,"text_chunks": text_chunks}

    except Exception as e:
        print(e)
        raise HTTPException(status_code=500, detail=str(e))





# GENERATE EMBEDDINGS FOR QUERY
class QueryEmbeddingRequest(BaseModel):
    text: str

@app.post("/generate-query-embedding")
def generate_query_embedding(request: QueryEmbeddingRequest):
    try:
        # Generate embedding for the input string (user's question)
        embedding = embedding_model.embed_query(request.text)
        
        return { "embedding": embedding }

    except Exception as e:
        print("Embedding error:", e)
        raise HTTPException(status_code=500, detail="Failed to generate embedding.")







class QueryRequest(BaseModel):
    question: str
    context: str
    previousChats: list[dict] = []

# You are an intelligent AI assistant. Answer the following question based on the provided context.
# ✅ Setup Groq Model
llm = ChatGroq(model_name="Gemma2-9b-It", groq_api_key=groq_api_key,temperature=0.7)

prompt_template = PromptTemplate(
    input_variables=["context", "question"],
    template="""
    You are an intelligent AI assistant. Consider the Previous Conversation (if present), then think **briefly** (at most 2 lines), and then provide a **long detailed response** (at least 3-4 lines) based on the given context.

    Format the response in valid HTML:
    - Wrap important information in `<strong>` tags.
    - Present key points using `<ul><li></li></ul>` lists when needed.
    - Use `<p>` tags to structure the response.

    Context:
    {context}

    Previous Conversations:
    {previousChats}

    Question: {question}

    Provide the direct answers only.
    """
)

# chain = llm | prompt_template
chain = LLMChain(llm=llm, prompt=prompt_template)

@app.post("/generate-response")
async def generate_response(request: QueryRequest):
    try:
        question = request.question
        context = request.context
        previousChats = request.previousChats

        formatted_chats = (
            "\n".join(
                [f"User: {chat['question']}\nAI: {chat['answer']}" for chat in previousChats]
            )
            if previousChats
            else "No previous conversation available."
        )


        # ✅ Use the Chain to Generate a Response
        response = chain.invoke({"context": context, "question": question,"previousChats": formatted_chats})
        # response = chain.invoke(formatted_input)

        return {"answer": response}

    except Exception as e:
        print(e)
        raise HTTPException(status_code=500, detail=str(e))





if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=5500)